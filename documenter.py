"""
Access DB Documenter
====================
Captures Design View + Datasheet View screenshots of every table in a
Microsoft Access .accdb / .mdb file, plus the Relationships diagram if
present, and writes them into a Word document.

Usage (GUI):   python documenter.py
Usage (CLI):   python documenter.py path/to/database.accdb
"""

import os
import sys
import time
import tempfile
import threading
import subprocess
import argparse
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path

# Must be set before any window / DPI calls
import ctypes
try:
    ctypes.windll.shcore.SetProcessDpiAwareness(1)   # PROCESS_SYSTEM_DPI_AWARE
except Exception:
    pass

import win32com.client
import win32gui
import win32con
import win32api
from PIL import Image, ImageGrab, ImageStat
from PIL import ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Access COM constants ───────────────────────────────────────────────────────
AC_VIEW_NORMAL       = 0    # Datasheet View
AC_VIEW_DESIGN       = 1    # Design View
AC_OBJ_TABLE         = 0    # AcObjectType.acTable
AC_SAVE_NONE         = 2    # AcCloseSave.acSaveNo
AC_QUIT_NO_SAVE      = 2    # AcQuitOption.acQuitSaveNone
AC_CMD_RELATIONSHIPS = 126  # Open Relationships window
AC_CMD_SHOW_ALL_REL  = 153  # Show all relationships in that window
AC_CMD_SHOW_ALL_REL_ALT = 149  # Alternate id seen in some Access builds


# ── Helpers ───────────────────────────────────────────────────────────────────

def _access_is_running() -> bool:
    result = subprocess.run(
        ["tasklist", "/fi", "imagename eq MSACCESS.EXE"],
        capture_output=True, text=True
    )
    return "MSACCESS.EXE" in result.stdout


def _find_access_hwnd():
    """Return the HWND of the main Access window (class OMain), or None."""
    found = []

    def _cb(hwnd, _):
        if win32gui.GetClassName(hwnd) == "OMain":
            found.append(hwnd)

    win32gui.EnumWindows(_cb, None)
    return found[0] if found else None


def _vk_press(vk, delay=0.15):
    """Press and release a single virtual key."""
    win32api.keybd_event(vk, 0, 0, 0)
    time.sleep(0.05)
    win32api.keybd_event(vk, 0, win32con.KEYEVENTF_KEYUP, 0)
    time.sleep(delay)


def _open_relationships_window(acc, hwnd, log=print, use_sendkeys=False):
    """
    Open the Access Relationships window.
    Primary path uses RunCommand(126), which maps to Database Tools ->
    Relationships and is language-independent.
    Then we try ShowAllRelationships (149) to ensure all links are visible.
    If 149 is unavailable, we keep the opened Relationships view as-is.
    """
    win32gui.ShowWindow(hwnd, win32con.SW_MAXIMIZE)
    win32gui.SetForegroundWindow(hwnd)
    time.sleep(0.4)

    # Primary: fire the ribbon control directly.
    # Some Access builds respond better to ExecuteMso than DoCmd.RunCommand.
    opened = False
    try:
        acc.CommandBars.ExecuteMso("Relationships")
        opened = True
        time.sleep(2.2)
    except Exception as ex:
        log(f"  ExecuteMso('Relationships') failed ({ex}); trying control-id 1629.")

    if not opened:
        try:
            ctrl = acc.CommandBars.FindControl(Id=1629)
            if ctrl is not None:
                ctrl.Execute()
                opened = True
                time.sleep(2.2)
            else:
                log("  CommandBars control 1629 not found; trying RunCommand fallback.")
        except Exception as ex:
            log(f"  CommandBars control 1629 failed ({ex}); trying RunCommand fallback.")

    if not opened:
        acc.DoCmd.RunCommand(AC_CMD_RELATIONSHIPS)
        time.sleep(2.2)

    # Maximize active object window to reduce clipping issues in Access MDI.
    try:
        acc.DoCmd.Maximize()
    except Exception:
        pass

    # Click into the document work area to ensure the Relationships tab is active.
    r = win32gui.GetWindowRect(hwnd)
    cx = (r[0] + r[2]) // 2
    cy = int(r[1] + (r[3] - r[1]) * 0.55)
    win32api.SetCursorPos((cx, cy))
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTDOWN, 0, 0)
    time.sleep(0.05)
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTUP, 0, 0)
    time.sleep(0.6)

    # Best effort: expand to all relationships if command is available.
    show_all_ok = False
    for _ in range(3):
        try:
            acc.DoCmd.RunCommand(AC_CMD_SHOW_ALL_REL)
            time.sleep(1.3)
            show_all_ok = True
            break
        except Exception:
            try:
                acc.DoCmd.RunCommand(AC_CMD_SHOW_ALL_REL_ALT)
                time.sleep(1.3)
                show_all_ok = True
                break
            except Exception:
                time.sleep(0.4)

    if show_all_ok:
        return True

    # Optional fallback: force ribbon focus/layout refresh with SendKeys.
    # Disabled by default because invalid keytips can trigger Windows beep sounds.
    if use_sendkeys:
        try:
            shell = win32com.client.Dispatch("WScript.Shell")
            shell.AppActivate(win32gui.GetWindowText(hwnd))
            time.sleep(0.5)
            shell.SendKeys("%Y2")
            time.sleep(0.6)
            shell.SendKeys("%JA")
            time.sleep(1.8)
            return True
        except Exception as ex:
            log(f"  SendKeys fallback failed ({ex}).")

    # Last fallback: temporary VBA helper + acc.Run().
    # This works on some builds where RunCommand ids are blocked.
    try:
        vbe = acc.VBE
        proj = vbe.ActiveVBProject
        comp = proj.VBComponents.Add(1)  # vbext_ct_StdModule
        comp.Name = "RelShowHelper"
        comp.CodeModule.InsertLines(
            1,
            "Public Sub ShowAllRelHelper()\n"
            "    DoCmd.ShowAllRelationships\n"
            "End Sub"
        )
        time.sleep(0.3)

        ran = False
        proj_name = ""
        try:
            proj_name = str(acc.CurrentProject.Name)
        except Exception:
            proj_name = ""

        attempts = [
            lambda: acc.Run("ShowAllRelHelper"),
            lambda: acc.Run("RelShowHelper.ShowAllRelHelper"),
            lambda: acc.Eval("ShowAllRelHelper()"),
        ]
        if proj_name:
            attempts.insert(1, lambda: acc.Run(f"{proj_name}.ShowAllRelHelper"))
            attempts.insert(2, lambda: acc.Run(f"{proj_name}.RelShowHelper.ShowAllRelHelper"))

        last_err = None
        for run_call in attempts:
            try:
                run_call()
                ran = True
                break
            except Exception as ex_run:
                last_err = ex_run

        if not ran and last_err is not None:
            raise last_err

        if ran:
            time.sleep(1.3)

        try:
            proj.VBComponents.Remove(comp)
        except Exception:
            pass
    except Exception as ex:
        log(f"  VBA ShowAllRelationships fallback failed ({ex}).")

    log("  ShowAllRelationships command unavailable; capturing current relationships view.")
    time.sleep(1.0)
    return False


def _capture_hwnd(hwnd, delay: float = 0.8) -> "Image.Image":
    """Maximize window, bring to foreground, then grab the full window."""
    win32gui.ShowWindow(hwnd, win32con.SW_MAXIMIZE)
    win32gui.SetForegroundWindow(hwnd)
    time.sleep(delay)
    rect = win32gui.GetWindowRect(hwnd)
    return ImageGrab.grab(bbox=rect)


def _capture_content(hwnd, delay: float = 0.8) -> "Image.Image":
    """
    Capture and crop to the Access content area only, removing:
      - left  ~13 % : navigation pane
      - top   ~16 % : title bar + ribbon + document-tab bar
      - bottom ~3 % : status bar
    These percentages hold for Access on typical 1080 p / 1440 p screens.
    """
    img = _capture_hwnd(hwnd, delay=delay)
    w, h = img.size
    l = int(w * 0.13)
    t = int(h * 0.16)
    b = h - int(h * 0.03)
    return img.crop((l, t, w, b))


def _largest_child_rect(hwnd):
    """Return rect of the largest visible child window of hwnd, or None."""
    candidates = []

    def _cb(ch, _):
        try:
            if not win32gui.IsWindowVisible(ch):
                return
            l, t, r, b = win32gui.GetWindowRect(ch)
            area = max(0, r - l) * max(0, b - t)
            if area < 80000:
                return
            cls = win32gui.GetClassName(ch)
            candidates.append((area, (l, t, r, b), cls))
        except Exception:
            return

    win32gui.EnumChildWindows(hwnd, _cb, None)
    if not candidates:
        return None
    candidates.sort(key=lambda x: x[0], reverse=True)
    return candidates[0][1]


def _has_document_surface(hwnd) -> bool:
    """Return True when Access has created a large document child window."""
    ignored_classes = {
        "MDIClient",
        "MsoCommandBar",
        "MsoCommandBarDock",
        "MsoWorkPane",
        "NUIPane",
        "NetUIHWND",
        "NetUINativeHWNDHost",
    }
    found = False

    def _cb(ch, _):
        nonlocal found
        try:
            if found or not win32gui.IsWindowVisible(ch):
                return
            l, t, r, b = win32gui.GetWindowRect(ch)
            area = max(0, r - l) * max(0, b - t)
            if area < 80000:
                return
            cls = win32gui.GetClassName(ch)
            if cls not in ignored_classes:
                found = True
        except Exception:
            return

    win32gui.EnumChildWindows(hwnd, _cb, None)
    return found


def _wait_for_document_surface(hwnd, timeout: float = 20.0, poll: float = 0.25) -> bool:
    """Wait until Access creates a real document surface child window."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        if _has_document_surface(hwnd):
            return True
        time.sleep(poll)
    return False


def _prompt_manual_relationships(hwnd) -> bool:
    """Ask the user to open Relationships manually in Access."""
    win32gui.ShowWindow(hwnd, win32con.SW_MAXIMIZE)
    win32gui.SetForegroundWindow(hwnd)
    time.sleep(0.3)

    message = (
        "Please open Relationships in Access.\n\n"
        "In Access, click Database Tools > Relationships.\n"
        "When you can see the relationships diagram, click OK.\n\n"
        "Click Cancel to skip this step."
    )
    title = "Open Relationships"
    MB_OKCANCEL = 0x00000001
    MB_ICONINFORMATION = 0x00000040
    IDOK = 1
    result = ctypes.windll.user32.MessageBoxW(None, message, title, MB_OKCANCEL | MB_ICONINFORMATION)
    return result == IDOK


def _capture_relationships(hwnd, delay: float = 1.4) -> "Image.Image":
    """
    Capture Relationships view with a safer crop strategy.
    We compare a full-window shot vs. content-cropped shot and keep the one
    with higher visual variance (usually the actual diagram instead of a blank area).
    """
    full_img = _capture_hwnd(hwnd, delay=delay)
    content_img = _capture_content(hwnd, delay=0.0)
    child_rect = _largest_child_rect(hwnd)
    child_img = None
    if child_rect is not None:
        child_img = ImageGrab.grab(bbox=child_rect)

    def _variance_score(img: "Image.Image") -> float:
        stat = ImageStat.Stat(img.convert("L"))
        return float(stat.var[0])

    best = full_img
    best_score = _variance_score(full_img)

    content_score = _variance_score(content_img)
    if content_score > best_score:
        best = content_img
        best_score = content_score

    if child_img is not None:
        child_score = _variance_score(child_img)
        if child_score > best_score:
            best = child_img

    return best


def _variance_score(img: "Image.Image") -> float:
    stat = ImageStat.Stat(img.convert("L"))
    return float(stat.var[0])


def _schema_relationships_image(db, width=1800, height=1100) -> "Image.Image":
    """Build a readable relationships diagram from DAO metadata."""
    bg = (247, 248, 250)
    box = (255, 255, 255)
    border = (80, 95, 120)
    title_bg = (38, 72, 117)
    title_fg = (255, 255, 255)
    line_col = (35, 35, 35)

    img = Image.new("RGB", (width, height), bg)
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    tables = _user_tables(db)
    rels = []
    for i in range(db.Relations.Count):
        rel = db.Relations(i)
        if rel.Name.startswith("MSys"):
            continue
        rels.append((rel.Table, rel.ForeignTable, rel.Name))

    if not tables:
        draw.text((30, 30), "No user tables found.", fill=(60, 60, 60), font=font)
        return img

    cols = max(2, min(4, (len(tables) + 2) // 3))
    box_w = max(320, (width - 80 - (cols - 1) * 40) // cols)
    x0 = 40
    y0 = 70
    col_gap = 40
    row_gap = 28
    title_h = 28
    field_h = 18

    table_fields = {}
    for t in tables:
        fields = []
        try:
            tdef = db.TableDefs(t)
            for fi in range(tdef.Fields.Count):
                fields.append(tdef.Fields(fi).Name)
        except Exception:
            pass
        table_fields[t] = fields[:18]

    table_rects = {}
    col_heights = [y0] * cols

    for idx, t in enumerate(tables):
        col = idx % cols
        tx = x0 + col * (box_w + col_gap)
        fields = table_fields[t]
        box_h = title_h + max(1, len(fields)) * field_h + 12
        ty = col_heights[col]

        draw.rectangle((tx, ty, tx + box_w, ty + box_h), fill=box, outline=border, width=2)
        draw.rectangle((tx, ty, tx + box_w, ty + title_h), fill=title_bg, outline=title_bg)
        draw.text((tx + 8, ty + 7), t, fill=title_fg, font=font)

        yy = ty + title_h + 6
        if fields:
            for f in fields:
                draw.text((tx + 10, yy), f"- {f}", fill=(35, 35, 35), font=font)
                yy += field_h
        else:
            draw.text((tx + 10, yy), "- (no fields)", fill=(80, 80, 80), font=font)

        table_rects[t] = (tx, ty, tx + box_w, ty + box_h)
        col_heights[col] += box_h + row_gap

    for t1, t2, rname in rels:
        if t1 not in table_rects or t2 not in table_rects:
            continue
        a = table_rects[t1]
        b = table_rects[t2]
        ax, ay = a[2], (a[1] + a[3]) // 2
        bx, by = b[0], (b[1] + b[3]) // 2
        if ax > bx:
            ax, bx = a[0], b[2]
        midx = (ax + bx) // 2
        draw.line((ax, ay, midx, ay, midx, by, bx, by), fill=line_col, width=2)
        draw.ellipse((ax - 3, ay - 3, ax + 3, ay + 3), fill=line_col)
        draw.ellipse((bx - 3, by - 3, bx + 3, by + 3), fill=line_col)
        draw.text((midx + 4, min(ay, by) + 2), rname, fill=(70, 70, 70), font=font)

    draw.text((40, 24), "Relationships (schema fallback)", fill=(20, 20, 20), font=font)
    return img



def _user_tables(db) -> list:
    """Return sorted list of user-visible table names (no MSys / temp)."""
    names = []
    for i in range(db.TableDefs.Count):
        name = db.TableDefs(i).Name
        if not name.startswith(("MSys", "USys", "~")):
            names.append(name)
    return sorted(names)


def _has_user_relations(db) -> bool:
    """True if the DB contains at least one non-system relationship."""
    for i in range(db.Relations.Count):
        if not db.Relations(i).Name.startswith("MSys"):
            return True
    return False


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Core logic ────────────────────────────────────────────────────────────────

def document_database(
    db_path: str,
    log=None,
    relationships_only: bool = False,
    relationships_shot_only: bool = False,
    use_sendkeys: bool = False,
) -> str:
    """
    Open *db_path* in Microsoft Access, capture screenshots, write .docx.
    Returns the path to the generated Word document.
    """
    if log is None:
        log = print

    db_path  = str(Path(db_path).resolve())
    out_path = str(Path(db_path).with_suffix(".docx"))
    rel_out_path = str(Path(db_path).with_name(f"{Path(db_path).stem}_relationships.png"))
    tmp_dir  = tempfile.mkdtemp(prefix="accdb_doc_")
    tmp_files: list = []

    log("Opening Microsoft Access…")
    acc = win32com.client.Dispatch("Access.Application")
    time.sleep(0.5)   # let COM server initialise before issuing any property set
    acc.OpenCurrentDatabase(db_path)
    acc.Visible = True
    time.sleep(2.5)   # let Access fully render

    hwnd = _find_access_hwnd()
    if not hwnd:
        acc.Quit(AC_QUIT_NO_SAVE)
        raise RuntimeError(
            "Could not find the Access window.\n"
            "Make sure Microsoft Access is installed."
        )

    win32gui.ShowWindow(hwnd, win32con.SW_MAXIMIZE)
    win32gui.SetForegroundWindow(hwnd)
    time.sleep(0.5)

    db     = acc.CurrentDb()
    tables = _user_tables(db)
    log(f"Found {len(tables)} table(s): {', '.join(tables)}")

    # ── Build Word document (unless PNG-only test mode) ───────────────────
    doc = None
    if not relationships_shot_only:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

    try:
        if not relationships_only and not relationships_shot_only:
            for idx, tbl in enumerate(tables, 1):
                log(f"[{idx}/{len(tables)}] {tbl}")

                # Design View ──────────────────────────────────────────────
                log("  → Design View")
                acc.DoCmd.OpenTable(tbl, AC_VIEW_DESIGN)
                design_path = os.path.join(tmp_dir, f"{tbl}_design.png")
                _capture_content(hwnd, delay=1.5).save(design_path)
                tmp_files.append(design_path)
                acc.DoCmd.Close(AC_OBJ_TABLE, tbl, AC_SAVE_NONE)
                time.sleep(0.4)

                # Datasheet View ───────────────────────────────────────────
                log("  → Datasheet View")
                acc.DoCmd.OpenTable(tbl, AC_VIEW_NORMAL)
                datasheet_path = os.path.join(tmp_dir, f"{tbl}_datasheet.png")
                _capture_content(hwnd, delay=1.5).save(datasheet_path)
                tmp_files.append(datasheet_path)
                acc.DoCmd.Close(AC_OBJ_TABLE, tbl, AC_SAVE_NONE)
                time.sleep(0.4)

                # Word section ─────────────────────────────────────────────
                doc.add_heading(tbl, level=1)

                ph = doc.add_paragraph(
                    "[Description — describe the purpose of this table here]"
                )
                ph.runs[0].italic = True
                ph.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

                doc.add_heading("Design View", level=2)
                doc.add_picture(design_path, width=Inches(6.0))

                doc.add_heading("Datasheet View", level=2)
                doc.add_picture(datasheet_path, width=Inches(6.0))

                _add_page_break(doc)

        # Relationships ────────────────────────────────────────────────────
        if _has_user_relations(db):
            log("Capturing Relationships diagram…")
            try:
                rel_window_populated = _open_relationships_window(
                    acc, hwnd, log, use_sendkeys=use_sendkeys
                )
                if not _has_document_surface(hwnd):
                    log("  Automated open did not create the Relationships designer surface.")
                    if _prompt_manual_relationships(hwnd):
                        if not _wait_for_document_surface(hwnd, timeout=20.0):
                            raise RuntimeError(
                                "Relationships was not detected after manual prompt. "
                                "Open it in Access and keep the actual relationships canvas visible."
                            )
                        rel_window_populated = True
                    else:
                        raise RuntimeError(
                            "Relationships screenshot cancelled because Access did not create the designer surface automatically."
                        )
                rel_path = os.path.join(tmp_dir, "relationships.png")
                rel_img = _capture_relationships(hwnd, delay=1.8)

                # If result looks too flat, retry once after re-opening relationships.
                if _variance_score(rel_img) < 120.0:
                    log("  Relationships image looked empty; retrying capture once…")
                    _open_relationships_window(acc, hwnd, log, use_sendkeys=use_sendkeys)
                    rel_img = _capture_relationships(hwnd, delay=2.2)

                if not rel_window_populated:
                    log("  Relationships window may not be fully populated; saving raw UI capture.")

                rel_img.save(rel_path)
                tmp_files.append(rel_path)

                if relationships_only or relationships_shot_only:
                    rel_img.save(rel_out_path)
                    log(f"Saved relationships image → {rel_out_path}")

                    if relationships_shot_only:
                        # Save debug variants to quickly identify wrong crop/focus.
                        dbg_full = str(Path(db_path).with_name(f"{Path(db_path).stem}_relationships_full.png"))
                        dbg_content = str(Path(db_path).with_name(f"{Path(db_path).stem}_relationships_content.png"))
                        _capture_hwnd(hwnd, delay=0.0).save(dbg_full)
                        _capture_content(hwnd, delay=0.0).save(dbg_content)
                        log(f"Saved debug full image → {dbg_full}")
                        log(f"Saved debug content image → {dbg_content}")

                if relationships_shot_only:
                    log("Skipping Word generation (relationships screenshot test mode).")
                    return rel_out_path

                doc.add_heading("Relationships", level=1)
                ph2 = doc.add_paragraph(
                    "[Description — describe the relationships between tables here]"
                )
                ph2.runs[0].italic = True
                ph2.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                doc.add_picture(rel_path, width=Inches(6.0))
            except Exception as exc:
                log(f"  (Relationships screenshot skipped: {exc})")

                if relationships_shot_only:
                    raise

        if relationships_shot_only:
            raise RuntimeError("No user relationships found in database.")

        log("Saving Word document…")
        try:
            doc.save(out_path)
        except PermissionError:
            # Output file is open in Word — save alongside it with a timestamp
            from datetime import datetime
            stem = Path(out_path).stem
            ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
            out_path = str(Path(out_path).with_name(f"{stem}_{ts}.docx"))
            doc.save(out_path)
            log("  (previous file was open — saved to a new name)")
        log(f"Done! → {out_path}")
        return out_path

    finally:
        try:
            acc.Quit(AC_QUIT_NO_SAVE)
        except Exception:
            pass
        for f in tmp_files:
            try:
                os.remove(f)
            except Exception:
                pass
        try:
            os.rmdir(tmp_dir)
        except Exception:
            pass


# ── GUI ───────────────────────────────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Access DB Documenter")
        self.resizable(False, False)
        self._build_ui()

    def _build_ui(self):
        f = ttk.Frame(self, padding=20)
        f.pack(fill=tk.BOTH, expand=True)

        ttk.Label(
            f, text="Access DB Documenter", font=("Segoe UI", 14, "bold")
        ).pack()
        ttk.Label(
            f,
            text=(
                "Select an .accdb file and click Generate.\n"
                "Microsoft Access will open automatically — don't click anything in it."
            ),
            justify=tk.CENTER,
        ).pack(pady=(4, 16))

        row = ttk.Frame(f)
        row.pack(fill=tk.X)
        self._path_var = tk.StringVar()
        ttk.Entry(row, textvariable=self._path_var, width=44).pack(
            side=tk.LEFT, fill=tk.X, expand=True
        )
        ttk.Button(row, text="Browse…", command=self._browse).pack(
            side=tk.LEFT, padx=(6, 0)
        )

        self._btn = ttk.Button(
            f, text="Generate Documentation", command=self._run
        )
        self._btn.pack(pady=12)

        self._bar = ttk.Progressbar(f, mode="indeterminate", length=420)
        self._bar.pack(fill=tk.X)

        self._status = tk.StringVar(value="Ready.")
        ttk.Label(
            f, textvariable=self._status, wraplength=460, justify=tk.LEFT
        ).pack(pady=(6, 0))

    def _browse(self):
        p = filedialog.askopenfilename(
            title="Select Access Database",
            filetypes=[
                ("Access Database", "*.accdb *.mdb"),
                ("All files", "*.*"),
            ],
        )
        if p:
            self._path_var.set(p)

    def _run(self):
        path = self._path_var.get().strip()
        if not path:
            messagebox.showwarning("No file selected", "Please select an .accdb file.")
            return
        if not os.path.exists(path):
            messagebox.showerror("File not found", f"File not found:\n{path}")
            return
        if _access_is_running():
            if not messagebox.askyesno(
                "Access is already open",
                "Microsoft Access is currently running.\n\n"
                "Please save and close it before continuing, "
                "then click Yes to proceed.",
            ):
                return

        self._btn.config(state=tk.DISABLED)
        self._bar.start(10)
        self._status.set("Starting…")

        def _worker():
            try:
                out = document_database(
                    path,
                    log=lambda m: self.after(0, self._status.set, m),
                )
                self.after(0, self._finish, out, None)
            except Exception as exc:
                self.after(0, self._finish, None, exc)

        threading.Thread(target=_worker, daemon=True).start()

    def _finish(self, out_path, err):
        self._bar.stop()
        self._btn.config(state=tk.NORMAL)
        if err:
            self._status.set(f"Error: {err}")
            messagebox.showerror("Error", str(err))
        else:
            self._status.set(f"Saved → {out_path}")
            if messagebox.askyesno(
                "Done!",
                f"Documentation saved to:\n{out_path}\n\nOpen the folder?",
            ):
                os.startfile(os.path.dirname(out_path))


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) > 1:
        parser = argparse.ArgumentParser(description="Access DB Documenter")
        parser.add_argument("db_path", help="Path to .accdb/.mdb file")
        parser.add_argument(
            "--relationships-only",
            action="store_true",
            help="Capture only the Relationships diagram (faster test mode)",
        )
        parser.add_argument(
            "--relationships-shot-only",
            action="store_true",
            help="Capture only Relationships and save a PNG (no Word generation)",
        )
        parser.add_argument(
            "--use-sendkeys",
            action="store_true",
            help="Enable SendKeys ribbon fallback (may trigger Windows dings)",
        )
        args = parser.parse_args()
        document_database(
            args.db_path,
            relationships_only=args.relationships_only,
            relationships_shot_only=args.relationships_shot_only,
            use_sendkeys=args.use_sendkeys,
        )
    else:
        App().mainloop()
